import sys
from PyQt4 import QtCore, QtGui
from form import Ui_Wizard
from docx import Document 
 
class MyDialog(QtGui.QWizard):
    def __init__(self, parent=None):
        QtGui.QWidget.__init__(self, parent)
        self.ui = Ui_Wizard()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(self.makeResume)
        self.ui.pushButton_2.clicked.connect(self.makeSmartResume)

    def makeResume(self):
        document = Document()

        document.add_heading(str(self.ui.name.text()),0)
        p = document.add_paragraph(
            str(self.ui.address.text())
            )
        p.add_run(' | ')
        p.add_run(
            str(self.ui.phoneNumber.text())
            )
        p.add_run(' | ')
        p.add_run(
            str(self.ui.email.text())
            )
        if str(self.ui.websites.toPlainText()) != "":
            p.add_run(' | ')
            p.add_run(str(self.ui.websites.toPlainText()))
        document.add_heading("Objective", 1)
        p2 = document.add_paragraph(
            str(self.ui.objective.toPlainText())
            )
        document.add_heading("Skills and Technologies", 1)
        skillsPar = document.add_paragraph(
            str(self.ui.strongSkills.toPlainText()), style="ListBullet"
        )
        document.add_paragraph(
            str(self.ui.strongTech.toPlainText()), style="ListBullet"
        )
        document.add_heading("Experience", 1)
        if(str(self.ui.jobEnd1.text()) == ""):
            jobEnd1 = "Present"
        else:
            jobEnd1 = str(self.ui.jobEnd1.text())
        document.add_heading(str(self.ui.jobStart1.text()) + " - " +
                             jobEnd1, 2)
        p3 = document.add_paragraph(
            str(self.ui.jobTitle1.text())
        )
        p3.add_run(", ")
        p3.add_run(str(self.ui.orgName1.text())).italic = True
        document.add_paragraph(
            str(self.ui.jobDesc1.toPlainText())
            )
        document.add_heading("Personal Development", 1)
        document.add_paragraph(
            str(self.ui.personalDev1.toPlainText())
            )
        document.add_heading('Achievements', 1)
        document.add_paragraph(
            str(self.ui.achievements.toPlainText())
            )
        document.add_heading('References', 1)
        document.add_paragraph(
            str(self.ui.references.toPlainText())
            )
        document.add_heading('Education', 1)
        if str(self.ui.endSchool1.text()) != "":
            document.add_heading(str(self.ui.schoolStart1.text()) + " - "
                                 + str(self.ui.endSchool1.text()), 2)
        else:
            document.add_heading(str(self.ui.schoolStart1.text()), 2)
        p4 = document.add_paragraph(str(self.ui.major1.text()))
        p4.add_run(", ")
        p4.add_run(str(self.ui.school1.text())).italic = True
        if self.ui.gpa1.value() != 0:
            p4.add_run(", Current GPA: " + str(self.ui.gpa1.value()))
        document.add_page_break()
        document.save('resume.docx')
        with open("dictionary.txt", "w") as f:
            f.write(str(self.ui.altSkills.toPlainText()))
        dicList = str(self.ui.altSkills.toPlainText()).split()
        jobDescList = str(self.ui.jobQualifications.toPlainText()).split()

    def makeSmartResume(self):
        document = Document()

        document.add_heading(str(self.ui.name.text()),0)
        p = document.add_paragraph(
            str(self.ui.address.text())
        )
        p.add_run(' | ')
        p.add_run(
            str(self.ui.phoneNumber.text())
        )
        p.add_run(' | ')
        p.add_run(
            str(self.ui.email.text())
        )
        if str(self.ui.websites.toPlainText()) != "":
            p.add_run(' | ')
            p.add_run(str(self.ui.websites.toPlainText()))
        document.add_heading("Objective", 1)
        p2 = document.add_paragraph(
            str(self.ui.objective.toPlainText())
        )
        document.add_heading("Skills and Technologies", 1)
        skillsPar = document.add_paragraph(
            str(self.ui.strongSkills.toPlainText()), style="ListBullet"
        )
        document.add_paragraph(
            str(self.ui.strongTech.toPlainText()), style="ListBullet"
        )
        document.add_heading("Experience", 1)
        if(str(self.ui.jobEnd1.text()) == ""):
            jobEnd1 = "Present"
        else:
            jobEnd1 = str(self.ui.jobEnd1.text())
            document.add_heading(str(self.ui.jobStart1.text()) + " - " +
                                 jobEnd1, 2)
        p3 = document.add_paragraph(
            str(self.ui.jobTitle1.text())
        )
        p3.add_run(", ")
        p3.add_run(str(self.ui.orgName1.text())).italic = True
        document.add_paragraph(
            str(self.ui.jobDesc1.toPlainText())
        )
        document.add_heading("Personal Development", 1)
        document.add_paragraph(
            str(self.ui.personalDev1.toPlainText())
        )
        document.add_heading('Achievements', 1)
        document.add_paragraph(
            str(self.ui.achievements.toPlainText())
        )
        document.add_heading('References', 1)
        document.add_paragraph(
            str(self.ui.references.toPlainText())
        )
        document.add_heading('Education', 1)
        if str(self.ui.endSchool1.text()) != "":
            document.add_heading(str(self.ui.schoolStart1.text()) + " - "
                                 + str(self.ui.endSchool1.text()), 2)
        else:
            document.add_heading(str(self.ui.schoolStart1.text()), 2)
        p4 = document.add_paragraph(str(self.ui.major1.text()))
        p4.add_run(", ")
        p4.add_run(str(self.ui.school1.text())).italic = True
        if self.ui.gpa1.value() != 0:
            p4.add_run(", Current GPA: " + str(self.ui.gpa1.value()))
        with open("dictionary.txt", "w") as f:
            f.write(str(self.ui.altSkills.toPlainText()))
            dicList = str(self.ui.altSkills.toPlainText()).lower().split()
            jobDescList = str(self.ui.jobQualifications.toPlainText()).lower().split()
        for descWords in jobDescList:
            for dicWords in dicList:
                if descWords.strip(".,") == dicWords:
                    skillsPar.add_run(", " + descWords.strip(".,"))
        document.add_page_break()
        document.save('resume.docx')

    
if __name__ == "__main__":
    app = QtGui.QApplication(sys.argv)
    myapp = MyDialog()
    myapp.show()
    sys.exit(app.exec_())


